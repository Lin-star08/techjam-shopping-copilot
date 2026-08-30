from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "_team_plan_assets"
OUT_PATH = OUT_DIR / "TechJam_Shopping_Copilot_5人团队三天作战手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
BLACK = RGBColor(0, 0, 0)
MUTED_FILL = "E8EEF5"
LIGHT_FILL = "F2F4F7"
CALL_FILL = "F4F6F9"
GREEN_FILL = "EAF6EF"
GOLD_FILL = "FFF4D6"
RED_FILL = "FCE8E6"
WHITE = "FFFFFF"


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def get_font(size: int = 34, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/Deng.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _wrap_one_line(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    tokens: list[str] = []
    current_ascii = ""
    for ch in text:
        if ch.isascii() and (ch.isalnum() or ch in "_@./+-"):
            current_ascii += ch
            continue
        if current_ascii:
            tokens.append(current_ascii)
            current_ascii = ""
        tokens.append(ch)
    if current_ascii:
        tokens.append(current_ascii)

    lines: list[str] = []
    current = ""
    for token in tokens:
        test = current + token
        width = draw.textbbox((0, 0), test, font=font)[2]
        if width <= max_width or not current:
            current = test
        else:
            lines.append(current.rstrip())
            current = token.lstrip()
    if current:
        lines.append(current.rstrip())
    return lines


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in text.splitlines():
        if not raw_line:
            lines.append("")
        else:
            lines.extend(_wrap_one_line(draw, raw_line, font, max_width))
    return lines


def rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    fill: str,
    outline: str,
    title: str,
    body: str,
    title_font: ImageFont.FreeTypeFont,
    body_font: ImageFont.FreeTypeFont,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    draw.text((x1 + 22, y1 + 18), title, font=title_font, fill="#0B2545")
    body_lines = wrap_text(draw, body, body_font, x2 - x1 - 44)
    y = y1 + 64
    for line in body_lines[:3]:
        draw.text((x1 + 22, y), line, font=body_font, fill="#24364B")
        y += 32


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#6B7A90", width=4)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        points = [(ex, ey), (ex - 16, ey - 10), (ex - 16, ey + 10)]
    else:
        points = [(ex, ey), (ex + 16, ey - 10), (ex + 16, ey + 10)]
    draw.polygon(points, fill="#6B7A90")


def create_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, True)
    box_title = get_font(25, True)
    body_font = get_font(22, False)
    small_font = get_font(19, False)
    draw.text((55, 38), "整体架构：数据驱动 + 稳健检索 + 多策略融合", font=title_font, fill="#0B2545")
    draw.text((58, 90), "目标：每轮给出 Top 10 商品推荐，同时问一个最有信息量的问题，尽早命中隐藏目标商品。", font=body_font, fill="#4B5563")

    boxes = [
        ((70, 165, 350, 330), "#E8EEF5", "1. 商品知识层", "品类规律、词典、失败案例到规则"),
        ((410, 165, 690, 330), "#EAF6EF", "2. 对话状态", "记住类别、材质、功能、已问属性"),
        ((750, 165, 1030, 330), "#FFF4D6", "3. 多路召回", "当前消息、历史、类别、属性、画像"),
        ((1090, 165, 1370, 330), "#F2F4F7", "4. 融合排序", "多路线加分、规则重排、Top 10"),
        ((1430, 165, 1710, 330), "#E8EEF5", "5. 回复策略", "推荐商品 + ask_attribute 追问"),
    ]
    for xy, fill, title, body in boxes:
        rounded_box(draw, xy, fill, "#B8C7D9", title, body, box_title, small_font)
    for x in [350, 690, 1030, 1370]:
        arrow(draw, (x + 10, 248), (x + 50, 248))

    rounded_box(
        draw,
        (245, 430, 705, 620),
        "#F7FAFC",
        "#B8C7D9",
        "输入数据",
        "5万商品目录 + 200条公开会话 + user_profile",
        box_title,
        body_font,
    )
    rounded_box(
        draw,
        (760, 430, 1220, 620),
        "#F7FAFC",
        "#B8C7D9",
        "本地评估器",
        "HitRate@10、MRR、MTTC、分场景结果",
        box_title,
        body_font,
    )
    rounded_box(
        draw,
        (1275, 430, 1650, 620),
        "#FCE8E6",
        "#E7A39D",
        "每日复盘",
        "看失败样本，把洞察变成下一轮规则",
        box_title,
        body_font,
    )
    arrow(draw, (1220, 525), (1275, 525))
    arrow(draw, (1450, 430), (1500, 345))
    arrow(draw, (480, 430), (210, 335))

    draw.text((58, 690), "原则：非技术成员负责让系统懂商品、懂用户和懂错误；技术成员负责把这些规则写进程序并用分数验证。", font=body_font, fill="#0B2545")
    img.save(path)


def create_timeline_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 520), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, True)
    label_font = get_font(25, True)
    body_font = get_font(21, False)
    draw.text((55, 40), "三天节奏：先搭骨架，再冲分，最后收口", font=title_font, fill="#0B2545")
    days = [
        ("Day 1 搭骨架", "目标：超过 baseline\n产出：状态记忆、词典 v1、追问策略 v1、基础 rerank"),
        ("Day 2 冲分", "目标：重点提升 browsing / override\n产出：多路召回、融合排序、失败复盘、参数实验"),
        ("Day 3 收口", "目标：稳定可提交\n产出：最终 agent、README、Devpost、demo 视频脚本"),
    ]
    x = 80
    for i, (title, body) in enumerate(days):
        fill = ["#E8EEF5", "#EAF6EF", "#FFF4D6"][i]
        xy = (x, 140, x + 500, 390)
        draw.rounded_rectangle(xy, radius=20, fill=fill, outline="#B8C7D9", width=3)
        draw.text((x + 28, 168), title, font=label_font, fill="#0B2545")
        y = 220
        for line in body.split("\n"):
            for wrapped in wrap_text(draw, line, body_font, 440):
                draw.text((x + 30, y), wrapped, font=body_font, fill="#24364B")
                y += 34
        if i < 2:
            arrow(draw, (x + 520, 265), (x + 610, 265))
        x += 610
    draw.text((75, 440), "硬规则：第3天下午以后不大改架构，只修 bug、整理交付、复查运行命令。", font=body_font, fill="#9B1C1C")
    img.save(path)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def setup_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "TechJam Shopping Copilot | 5人团队三天作战手册"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, 9, GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = "内部协作版 | 每天用评估分数决定取舍"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, 9, GRAY)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_table(table, widths: list[int], header_fill: str = MUTED_FILL) -> None:
    set_table_geometry(table, widths)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, 9.5)
            if i == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, 9.5, NAVY, True)


def add_para(doc: Document, text: str, size: float = 11, color: RGBColor = BLACK, bold: bool = False, italic: bool = False, align: int | None = None, after: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    for run in p.runs:
        set_run_font(run, 10.5)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    for run in p.runs:
        set_run_font(run, 10.5)


def add_callout(doc: Document, title: str, text: str, fill: str = CALL_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, 10.5, NAVY, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, 10.5, BLACK)
    shade_cell(cell, fill)
    format_table(table, [9360], fill)
    doc.add_paragraph()


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths_in: list[float], header_fill: str = MUTED_FILL) -> None:
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value
    format_table(table, [dxa(x) for x in widths_in], header_fill)
    doc.add_paragraph()


def add_title_page(doc: Document) -> None:
    add_para(doc, "TechJam Shopping Copilot", 12, GRAY, True, after=4)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("5人团队三天作战手册")
    set_run_font(r, 28, NAVY, True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r2 = subtitle.add_run("方案组合：稳健冲分型 + 数据驱动型 + 多策略融合型")
    set_run_font(r2, 14, GRAY)
    add_simple_table(
        doc,
        ["项目", "说明"],
        [
            ["比赛目标", "在最多10轮对话内，让隐藏目标商品尽早进入推荐 Top 10，并尽量排在前面。"],
            ["团队配置", "5人小组；技术和非技术成员都承担核心模块，不把非技术同学降级为打杂。"],
            ["主线策略", "离线可复现 agent：多轮记忆、商品知识层、多路召回、融合排序、主动追问。"],
            ["工作原则", "每个任务都要有可交付成果；每个算法改动都用 HitRate、MRR、MTTC 验证。"],
            ["文档日期", "2026-08-29"],
        ],
        [1.25, 5.25],
    )
    add_callout(
        doc,
        "组长开场话术",
        "我们不是做购物网站，而是做一个购物推荐大脑。非技术成员负责把商品、用户和错误案例讲清楚；技术成员负责把这些规则写成程序。最终所有想法都要接受分数检验。",
        GREEN_FILL,
    )


def build_doc() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    create_architecture_diagram(ASSET_DIR / "architecture.png")
    create_timeline_diagram(ASSET_DIR / "timeline.png")

    doc = Document()
    setup_sections(doc)
    setup_styles(doc)
    add_title_page(doc)

    add_heading(doc, "1. 最终目标与核心打法", 1)
    add_para(doc, "我们采用三套方案叠加，但不是做三套系统，而是把它们合成一条清晰主线。")
    add_simple_table(
        doc,
        ["方案", "在项目里的真实含义", "最终进入系统的产出"],
        [
            ["稳健冲分型", "先把官方 baseline 升级成稳定、离线、可复现的强 agent。", "session state、槽位抽取、BM25 多路检索、rerank、ask_attribute。"],
            ["数据驱动型", "先看数据和失败案例，再把观察转成规则。", "商品知识层、词典、品类 playbook、失败样本到规则的转化。"],
            ["多策略融合型", "不靠单一路线决定推荐，而是多路搜索后融合排序。", "当前消息、历史需求、类别、属性、用户画像五路召回。"],
        ],
        [1.35, 2.55, 2.60],
    )
    add_callout(doc, "一句话版本", "做一个会记住需求、会问问题、会从多个角度搜索商品，并用本地评估器持续变强的购物推荐 Agent。", GOLD_FILL)

    add_heading(doc, "2. 整体架构", 1)
    add_para(doc, "系统每一轮都执行同一个闭环：理解当前消息，更新会话状态，多路搜索商品，融合排序，返回 Top 10，并问一个最有价值的问题。")
    doc.add_picture(str(ASSET_DIR / "architecture.png"), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_simple_table(
        doc,
        ["模块", "负责解决的问题", "主要产出"],
        [
            ["商品知识层", "商品类别、材质、功能、品牌这些规律不能只靠程序猜。", "词典、品类规则、失败案例改进建议。"],
            ["对话状态层", "用户是多轮表达需求，系统必须记住有效信息，并处理改变主意。", "SessionState、slots、已问属性、无偏好属性。"],
            ["多路召回层", "单一路线容易漏掉目标商品，所以先从多个角度捞候选。", "5路候选列表，每路 Top 50 或 Top 100。"],
            ["融合排序层", "候选商品很多，必须把更像目标的商品排到前面。", "融合分数、场景权重、最终 Top 10。"],
            ["评估复盘层", "不靠感觉判断好坏，用公开集分数和失败样本迭代。", "results.json、实验记录、下一轮规则。"],
        ],
        [1.35, 3.05, 2.10],
    )

    add_heading(doc, "3. 5个人的角色分工", 1)
    add_para(doc, "每个人都负责一个正式模块。任务写得简单清楚，但每个模块都要能影响最终分数。")
    add_simple_table(
        doc,
        ["编号", "角色", "一句话职责", "主要交付物"],
        [
            ["1号", "数据洞察与商品知识层负责人", "把数据观察变成程序能使用的商品规则。", "category_playbook.md、lexicon_rules.xlsx、failure_to_rule.md、daily_error_review.md"],
            ["2号", "对话策略与场景负责人", "设计什么时候问什么问题，以及用户改变主意时怎么办。", "question_policy.md、question_templates.md、scenario_rules.md"],
            ["3号", "检索工程负责人", "从5万商品里尽量多捞出可能正确的候选。", "stateful Agent、5路召回函数、候选合并接口"],
            ["4号", "融合排序与调参负责人", "把正确商品尽量排到 Top 10 靠前位置。", "rerank 函数、融合权重、参数实验结论"],
            ["5号", "评估实验与交付负责人", "记录每次改动是否真的变强，并负责最终提交材料。", "experiments.md、results.json、README、demo_script.md"],
        ],
        [0.55, 1.55, 2.35, 2.05],
    )

    add_heading(doc, "4. 1号任务卡：数据洞察与商品知识层", 1)
    add_callout(doc, "定位", "1号不是只做表的人，而是负责让系统具备“商品常识”。他/她的产出会直接进入检索词典、追问策略和排序规则。", GREEN_FILL)
    add_simple_table(
        doc,
        ["任务", "具体要做什么", "交给谁使用", "完成标准"],
        [
            ["目标商品画像", "看200个公开样本的正确商品，整理高频品类、品牌、材质、功能词。", "3号、4号", "每个重点品类至少写出典型词和推荐加分方向。"],
            ["商品词典", "整理材质、颜色、功能、风格、用途、品牌词，并标注适用品类。", "3号", "格式能被程序转换成 lexicon.py 或 JSON。"],
            ["失败到规则", "看 baseline 和新版本 miss 的样本，写出为什么错以及下一步应加什么规则。", "4号、组长", "每天至少产出10条可执行规则建议。"],
            ["品类追问建议", "告诉2号不同品类优先问什么，例如鞋问 use_case，内衣问 fit/style。", "2号", "每个重点品类都有前3个追问属性。"],
            ["每日复盘", "每晚看新增命中和仍未命中的样本，总结哪些规则有效。", "全组", "写出“保留/调整/放弃”的明确建议。"],
        ],
        [1.15, 2.75, 0.85, 1.75],
    )
    add_para(doc, "给1号的口头说明：你每天要交的不是漂亮表格，而是能让程序变聪明的规则。比如你发现 bras 常出现 wireless、stretch、nylon，排序工程师就会给这些词加分；你发现 shoes 更看用途，2号就会把 use_case 放到前面问。", 10.5)

    add_heading(doc, "5. 2号任务卡：对话策略与场景", 1)
    add_simple_table(
        doc,
        ["任务", "具体要做什么", "交给谁使用", "完成标准"],
        [
            ["追问顺序", "为 browsing、buying、intent_override、boundary 四类场景写追问优先级。", "3号", "能映射到 ask_attribute，且不会重复问用户说无偏好的属性。"],
            ["问题模板", "给 material、style、feature、brand、color、size、use_case、budget 写简洁英文问题。", "3号", "每个属性至少2个模板，语气自然。"],
            ["无偏好处理", "用户说 no preference 时，标记该属性为 neutral，下一轮换问别的。", "3号、4号", "有清楚规则，不让系统原地打转。"],
            ["意图覆盖处理", "用户说 Actually / ignore earlier 时，以新需求为准，旧需求降权或删除。", "3号、4号", "写出触发词、清理规则、下一轮动作。"],
            ["demo 对话", "准备3个可展示案例：buying、browsing、override。", "5号", "每个案例有用户话术、系统问题、推荐逻辑解释。"],
        ],
        [1.15, 2.75, 0.85, 1.75],
    )
    add_callout(doc, "给2号的口头说明", "你负责让程序像一个清醒的导购：信息少时问关键问题，用户没偏好时换方向，用户改变主意时马上改路线。", CALL_FILL)

    add_heading(doc, "6. 3号任务卡：检索工程", 1)
    add_simple_table(
        doc,
        ["检索路线", "要搜索什么", "输入", "输出"],
        [
            ["当前消息搜索", "只看用户这一轮刚说的话，快速响应新需求。", "user_message", "route_current: Top 50"],
            ["历史需求搜索", "把用户前几轮有效需求合起来，避免丢上下文。", "SessionState.history", "route_history: Top 100"],
            ["类别搜索", "重点匹配 Dresses、Bras、Belts、Sneakers 等目录类别。", "state.category_terms", "route_category: Top 100"],
            ["属性搜索", "重点匹配 leather、cotton、waterproof、wireless、stretch 等。", "state.slots", "route_attribute: Top 100"],
            ["用户画像搜索", "根据 fit、comfort、material 等画像补充同义词。", "user_profile", "route_profile: Top 50"],
        ],
        [1.2, 2.2, 1.35, 1.75],
    )
    add_para(doc, "3号需要写出的代码接口建议：", 10.5, bold=True)
    add_simple_table(
        doc,
        ["函数", "说明"],
        [
            ["extract_slots(user_message)", "从用户话里抽取类别、材质、颜色、品牌、功能等关键词。"],
            ["update_state(session_id, user_message)", "更新本轮会话状态，处理历史、已问属性、意图覆盖。"],
            ["retrieve_candidates(state)", "运行5路召回，把候选商品和来源路线一起返回。"],
            ["choose_next_question(state)", "根据2号规则决定下一轮 ask_attribute。"],
        ],
        [2.1, 4.4],
        LIGHT_FILL,
    )

    add_heading(doc, "7. 4号任务卡：融合排序与调参", 1)
    add_para(doc, "4号的目标是提高 MRR：不仅要命中，还要把正确商品排得尽量靠前。")
    add_simple_table(
        doc,
        ["加分项", "建议初始权重", "说明"],
        [
            ["类别匹配", "+5", "用户说 Bras，就优先让 Everyday Bras、Lingerie、Women 相关商品靠前。"],
            ["材质匹配", "+4", "leather、cotton、nylon、polyester 等词对 buying 很关键。"],
            ["功能词匹配", "+3", "waterproof、wireless、stretch、warm、breathable 等词通常很有区分度。"],
            ["品牌/store 匹配", "+2到+3", "如果用户或画像出现品牌，适度加分。"],
            ["多路线同时出现", "+2", "同一个商品被多条路线召回，说明更稳。"],
            ["评分与评论数", "+0.2到+0.8", "只能轻微加分，不能压过明确需求。"],
            ["旧意图惩罚", "-3", "用户说 ignore earlier 后，旧偏好命中的商品要降权。"],
        ],
        [1.6, 1.2, 3.7],
    )
    add_callout(doc, "调参原则", "每次只改一类权重，跑完整公开集，看总分和分场景分数。变好保留，变差撤回，变化不明显就看是否提升 browsing 或 override。", GOLD_FILL)

    add_heading(doc, "8. 5号任务卡：评估实验与交付", 1)
    add_simple_table(
        doc,
        ["任务", "具体要做什么", "完成标准"],
        [
            ["跑分记录", "每次代码改动后运行本地评估器，记录总分和分场景分数。", "experiments.md 中有版本、改动、分数、结论。"],
            ["实验看板", "维护 v0、v1、v2 等版本对比，避免团队凭感觉争论。", "任何人都能看出哪个版本最好。"],
            ["交付材料", "写 README、Devpost 描述、模型/成本/限制说明。", "第三天中午前有可提交草稿。"],
            ["Demo 脚本", "准备一个 buying、一个 browsing、一个 override 的演示路径。", "能解释系统为什么问、为什么推荐。"],
            ["最终检查", "确认没有 API key、没有私有数据、没有改 evaluator 来报分。", "提交前完成检查清单。"],
        ],
        [1.25, 3.2, 2.05],
    )
    add_para(doc, "实验记录模板：", 10.5, bold=True)
    add_simple_table(
        doc,
        ["版本", "改了什么", "HitRate", "MRR", "MTTC", "总分", "结论"],
        [
            ["v0", "官方 baseline", "0.125", "0.068", "9.81", "0.1067", "对照组"],
            ["v1", "加多轮状态和基础 rerank", "", "", "", "", "待验证"],
            ["v2", "加商品知识层和多路融合", "", "", "", "", "待验证"],
        ],
        [0.65, 2.2, 0.65, 0.55, 0.55, 0.6, 1.3],
        LIGHT_FILL,
    )

    add_heading(doc, "9. 三天执行计划", 1)
    doc.add_picture(str(ASSET_DIR / "timeline.png"), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_simple_table(
        doc,
        ["时间", "1号 商品知识层", "2号 对话策略", "3号 检索工程", "4号 排序融合", "5号 评估交付"],
        [
            ["Day 1 上午", "读公开样本，列目标品类 Top 20。", "整理四类场景该问什么。", "读 agent.py，设计 SessionState。", "读 results.json，确认 baseline 短板。", "建立 experiments.md 和版本命名。"],
            ["Day 1 下午", "交词典 v1：材质、颜色、功能、品牌。", "交问题模板 v1。", "实现历史搜索和类别搜索。", "实现基础加分规则。", "跑 v1，记录分场景结果。"],
            ["Day 1 晚上", "分析10个失败样本，写规则建议。", "补 no preference 和 ignore earlier 规则。", "合并进 agent.py。", "调第一轮权重。", "确定 Day 2 优先攻哪个场景。"],
            ["Day 2 上午", "看 v1 miss，标失败原因。", "优化 browsing 追问。", "实现属性搜索和画像搜索。", "加入路线融合分。", "跑 v2，做对比。"],
            ["Day 2 下午", "补重点品类 playbook。", "准备 demo 对话草稿。", "优化候选合并和兜底。", "调 buying/browsing/override 权重。", "每2小时更新实验表。"],
            ["Day 2 晚上", "总结有效规则。", "锁定最终追问策略。", "修明显 bug。", "冻结候选最终权重。", "选出候选最终版。"],
            ["Day 3 上午", "整理数据洞察给报告。", "整理场景策略说明。", "代码清理，保证可运行。", "最后小幅调参。", "写 README 和 Devpost。"],
            ["Day 3 下午", "准备展示解释。", "配合录 demo。", "只修 bug，不大改。", "只修 bug，不大改。", "最终跑分，录视频，提交检查。"],
        ],
        [0.8, 1.14, 1.14, 1.14, 1.14, 1.14],
    )

    add_heading(doc, "10. 协作规则与接口", 1)
    add_simple_table(
        doc,
        ["规则", "具体做法"],
        [
            ["每天两次同步", "上午15分钟说今天交什么；晚上30分钟看分数和失败案例。"],
            ["所有产出可接入", "词表、规则、代码、实验记录都要能被其他人直接使用。"],
            ["版本命名统一", "v0 baseline，v1 stateful，v2 fusion，v3 final-candidate；每个版本都留分数。"],
            ["不争论感觉", "想法可以大胆，但保留与否看 HitRate、MRR、MTTC 和分场景表现。"],
            ["第3天控风险", "第三天下午以后不重构，只修 bug、写文档、录 demo、检查提交。"],
        ],
        [1.45, 5.05],
    )
    add_para(doc, "技术接口建议：3号和4号约定候选格式，避免最后拼不起来。", 10.5, bold=True)
    add_simple_table(
        doc,
        ["字段", "含义"],
        [
            ["parent_asin", "商品 ID，最终推荐只看这个字段。"],
            ["route_scores", "不同路线给出的原始或归一化分数。"],
            ["matched_terms", "命中的类别、材质、功能、品牌词。"],
            ["final_score", "4号融合排序后的总分。"],
            ["debug_reason", "给5号看失败案例时使用，不一定提交给评估器。"],
        ],
        [1.4, 5.1],
        LIGHT_FILL,
    )

    add_heading(doc, "11. 最终验收清单", 1)
    checks = [
        "agent.py 可以用 python -m evaluator.local_evaluator 跑完整公开集。",
        "results.json 中总分和四类场景分数已经记录到 experiments.md。",
        "README 写清楚安装、运行、复现结果、模型/成本、限制和团队贡献。",
        "没有提交 API key、私有数据、无关大文件或修改后的 evaluator 报分。",
        "Demo 至少展示 buying、browsing、intent_override 三种情况之一，最好三种都有。",
        "每个人都能用一句话说明自己模块如何让系统变强。",
    ]
    for item in checks:
        add_bullet(doc, item)
    add_callout(
        doc,
        "组长最后提醒",
        "这套分工的关键不是让非技术成员做轻活，而是把复杂技术任务翻译成清楚的产出：商品知识、对话规则、失败复盘、实验证据。这些都会进入最终 agent。",
        GREEN_FILL,
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
